import re
import os
import docx
import pdfplumber
import markdown
from docx.shared import Pt
import pypandoc
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

# 文本格式化函数
def format_text_run(run):
    """格式化单个文本段落，处理粗体、斜体和下划线"""
    content = run.text
    if not content.strip():
        return ""
        
    if run.bold:
        content = f"**{content}**"
    if run.italic:
        content = f"*{content}*"
    if run.underline:
        content = f"<u>{content}</u>"
    return content

# 表格处理函数
def convert_table_to_md(table):
    """将Word表格转换为Markdown表格"""
    if not table.rows:
        return ""
        
    # 表头
    header_cells = table.rows[0].cells
    header_row = [cell.text.strip() or " " for cell in header_cells]
    
    # 构建markdown表格
    md_rows = [
        "| " + " | ".join(header_row) + " |",
        "| " + " | ".join(["---"] * len(header_row)) + " |"
    ]
    
    # 表格内容
    for row in table.rows[1:]:
        cells = [cell.text.strip() or " " for cell in row.cells]
        md_rows.append("| " + " | ".join(cells) + " |")
    
    return "\n".join(md_rows)

# 标题处理函数
def process_heading(para):
    """处理标题段落，返回markdown格式的标题"""
    if para.style.name.startswith('Heading'):
        level = int(para.style.name.replace('Heading ', ''))
        return '#' * level + ' ' + para.text
    return None

# 从PDF提取文本
def extract_text_from_pdf(pdf_path):
    """从PDF文件提取文本内容"""
    content = []
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                # 添加页码标记
                content.append(f"## 第{page_num + 1}页")
                
                # 提取文本
                text = page.extract_text()
                if text:
                    content.append(text)
                
                # 尝试提取表格
                tables = page.extract_tables()
                for table in tables:
                    # 转换为Markdown表格
                    md_table = []
                    
                    # 表头
                    if table and table[0]:
                        header_row = [cell or " " for cell in table[0]]
                        md_table.append("| " + " | ".join(header_row) + " |")
                        md_table.append("| " + " | ".join(["---"] * len(header_row)) + " |")
                    
                        # 表格内容
                        for row in table[1:]:
                            cells = [cell or " " for cell in row]
                            md_table.append("| " + " | ".join(cells) + " |")
                        
                        content.append("\n".join(md_table))
                
    except Exception as e:
        content.append(f"PDF处理错误: {str(e)}")
    
    return "\n\n".join(content)

# 将Markdown转换为Word文档
def convert_md_to_word(md_path, output_path):
    """
    将Markdown文件转换为Word文档
    :param md_path: Markdown文件路径
    :param output_path: 输出Word文件路径
    :return: 是否成功
    """
    try:
        # 使用pypandoc转换
        pypandoc.convert_file(
            md_path,
            'docx',
            outputfile=output_path,
            extra_args=['--reference-doc=reference.docx'] if os.path.exists('reference.docx') else []
        )
        return True
    except Exception as e:
        print(f"转换失败: {str(e)}")
        # 备用方法：使用python-docx手动转换
        try:
            # 读取Markdown文件
            with open(md_path, 'r', encoding='utf-8') as f:
                md_content = f.read()
            
            # 创建Word文档
            doc = docx.Document()
            
            # 处理标题和内容
            lines = md_content.split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    doc.add_paragraph()
                    continue
                    
                # 处理标题
                if line.startswith('#'):
                    level = 0
                    while line.startswith('#'):
                        level += 1
                        line = line[1:]
                    line = line.strip()
                    doc.add_heading(line, level=level)
                # 处理普通段落
                else:
                    p = doc.add_paragraph()
                    # 简单处理粗体和斜体
                    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', line)
                    for part in parts:
                        if part.startswith('**') and part.endswith('**'):
                            p.add_run(part[2:-2]).bold = True
                        elif part.startswith('*') and part.endswith('*'):
                            p.add_run(part[1:-1]).italic = True
                        elif part:
                            p.add_run(part)
            
            # 保存文档
            doc.save(output_path)
            return True
        except Exception as e2:
            print(f"备用方法也失败: {str(e2)}")
            return False

# 将Markdown转换为PDF文档
def convert_md_to_pdf(md_path, output_path):
    """
    将Markdown文件转换为PDF
    :param md_path: Markdown文件路径
    :param output_path: 输出PDF路径
    :return: 是否成功
    """
    try:
        # 尝试使用pypandoc
        pypandoc.convert_file(
            md_path,
            'pdf',
            outputfile=output_path
        )
        return True
    except Exception as e:
        print(f"pypandoc转换失败: {str(e)}")
        
        # 备用方法：使用reportlab手动转换
        try:
            # 读取Markdown文件
            with open(md_path, 'r', encoding='utf-8') as f:
                md_content = f.read()
            
            # 将Markdown转换为HTML
            html = markdown.markdown(md_content, extensions=['tables', 'fenced_code'])
            
            # 创建PDF文档
            doc = SimpleDocTemplate(output_path, pagesize=letter)
            styles = getSampleStyleSheet()
            
            # 自定义标题样式
            styles.add(ParagraphStyle(name='Heading1', fontSize=16, bold=True))
            styles.add(ParagraphStyle(name='Heading2', fontSize=14, bold=True))
            styles.add(ParagraphStyle(name='Heading3', fontSize=12, bold=True))
            
            # 解析HTML并添加到PDF
            flowables = []
            
            # 简单解析HTML
            lines = html.split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    flowables.append(Spacer(1, 12))
                    continue
                
                # 处理标题
                if line.startswith('<h1>'):
                    text = line.replace('<h1>', '').replace('</h1>', '')
                    flowables.append(Paragraph(text, styles['Heading1']))
                elif line.startswith('<h2>'):
                    text = line.replace('<h2>', '').replace('</h2>', '')
                    flowables.append(Paragraph(text, styles['Heading2']))
                elif line.startswith('<h3>'):
                    text = line.replace('<h3>', '').replace('</h3>', '')
                    flowables.append(Paragraph(text, styles['Heading3']))
                # 处理段落
                elif not (line.startswith('<table>') or line.startswith('<tr>') or line.startswith('<td>')):
                    # 移除其他HTML标签
                    text = re.sub(r'<[^>]*>', '', line)
                    flowables.append(Paragraph(text, styles['Normal']))
            
            # 构建PDF
            doc.build(flowables)
            return True
        except Exception as e2:
            print(f"备用方法也失败: {str(e2)}")
            return False

# 将多个Markdown文件合并为一个
def merge_markdown_files(md_paths, output_path):
    """
    合并多个Markdown文件为一个
    :param md_paths: Markdown文件路径列表
    :param output_path: 输出的合并文件路径
    """
    merged_content = []
    
    for md_path in md_paths:
        try:
            file_name = os.path.splitext(os.path.basename(md_path))[0]
            with open(md_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # 添加文件标题和内容
            merged_content.append(f"# {file_name}\n\n{content}")
        except Exception as e:
            print(f"合并文件 {md_path} 时出错: {str(e)}")
    
    # 写入合并后的文件
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write("\n\n---\n\n".join(merged_content))
    
    return True

# 新增函数: 简单模式Word文档转Markdown
def extract_text_simple(doc):
    """简单模式：提取Word文档的文本并保留格式"""
    content = []
    
    # 处理段落
    for para in doc.paragraphs:
        # 空段落
        if not para.text.strip():
            content.append("\n")
            continue
        
        # 处理标题
        if heading := process_heading(para):
            content.append(heading)
            continue
            
        # 处理普通段落
        formatted_runs = [format_text_run(run) for run in para.runs]
        formatted_text = "".join(formatted_runs)
        if formatted_text:
            content.append(formatted_text)
    
    # 处理表格
    for table in doc.tables:
        content.append(convert_table_to_md(table))
    
    return "\n\n".join(content)

# 新增函数: 分割模式Word文档转Markdown
def extract_text_with_sections(doc):
    """分割模式：按一级标题提取Word文档内容并分割为多个章节"""
    sections = {}
    current_section = []
    current_title = "前言"  # 默认标题
    
    # 处理段落
    for para in doc.paragraphs:
        # 处理一级标题 - 分割点
        if para.style.name == 'Heading 1':
            # 保存之前的部分
            if current_section:
                sections[current_title] = "\n\n".join(current_section)
            # 开始新的部分
            current_title = para.text.strip()
            current_section = []
            current_section.append('# ' + para.text)
            continue
            
        # 空段落
        if not para.text.strip():
            current_section.append("\n")
            continue
        
        # 处理其他级别标题
        if para.style.name.startswith('Heading'):
            level = int(para.style.name.replace('Heading ', ''))
            current_section.append('#' * level + ' ' + para.text)
            continue
            
        # 处理普通段落
        formatted_runs = [format_text_run(run) for run in para.runs]
        formatted_text = "".join(formatted_runs)
        if formatted_text:
            current_section.append(formatted_text)
    
    # 处理表格
    for table in doc.tables:
        current_section.append(convert_table_to_md(table))
    
    # 保存最后一个部分
    if current_section:
        sections[current_title] = "\n\n".join(current_section)
    
    return sections 