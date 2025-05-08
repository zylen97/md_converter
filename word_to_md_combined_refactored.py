import docx
import re
import os
import sys
import pdfplumber  # PDF处理库
import markdown  # Markdown处理
import docxtpl  # Word模板
from docx.shared import Pt, Inches  # Word文档样式
from reportlab.lib.pagesizes import letter  # PDF生成
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
import pypandoc  # 通用文档转换
from PyQt5.QtWidgets import (QApplication, QMainWindow, QPushButton, QLabel, 
                           QFileDialog, QVBoxLayout, QHBoxLayout, QWidget, 
                           QTextEdit, QProgressBar, QMessageBox, QListWidget,
                           QGroupBox, QRadioButton, QButtonGroup, QCheckBox, QTabWidget)
from PyQt5.QtCore import Qt, QThread, pyqtSignal
from PyQt5.QtGui import QFont

# 导入自定义模块
from converters import ToMarkdownThread, FromMarkdownThread
from utils import (format_text_run, convert_table_to_md, process_heading, 
                   extract_text_from_pdf, convert_md_to_word, convert_md_to_pdf,
                   extract_text_simple, extract_text_with_sections)

# 辅助函数：格式化文本运行
def format_text_run(run):
    """格式化单个文本段落，处理粗体、斜体和下划线"""
    content = run.text
    if not content.strip():
        return ""
        
    if run.bold:
        content = f"**{content}**"
    if run.italic:
        content = f"*{content}*"
    if run.underline:
        content = f"<u>{content}</u>"
    return content

# 辅助函数：处理表格
def convert_table_to_md(table):
    """将Word表格转换为Markdown表格"""
    if not table.rows:
        return ""
        
    # 表头
    header_cells = table.rows[0].cells
    header_row = [cell.text.strip() or " " for cell in header_cells]
    
    # 构建markdown表格
    md_rows = [
        "| " + " | ".join(header_row) + " |",
        "| " + " | ".join(["---"] * len(header_row)) + " |"
    ]
    
    # 表格内容
    for row in table.rows[1:]:
        cells = [cell.text.strip() or " " for cell in row.cells]
        md_rows.append("| " + " | ".join(cells) + " |")
    
    return "\n".join(md_rows)

# 辅助函数：获取标题级别和内容
def process_heading(para):
    """处理标题段落，返回markdown格式的标题"""
    if para.style.name.startswith('Heading'):
        level = int(para.style.name.replace('Heading ', ''))
        return '#' * level + ' ' + para.text
    return None

# 从PDF提取文本
def extract_text_from_pdf(pdf_path):
    """从PDF文件提取文本内容"""
    content = []
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                # 添加页码标记
                content.append(f"## 第{page_num + 1}页")
                
                # 提取文本
                text = page.extract_text()
                if text:
                    content.append(text)
                
                # 尝试提取表格
                tables = page.extract_tables()
                for table in tables:
                    # 转换为Markdown表格
                    md_table = []
                    
                    # 表头
                    if table and table[0]:
                        header_row = [cell or " " for cell in table[0]]
                        md_table.append("| " + " | ".join(header_row) + " |")
                        md_table.append("| " + " | ".join(["---"] * len(header_row)) + " |")
                    
                        # 表格内容
                        for row in table[1:]:
                            cells = [cell or " " for cell in row]
                            md_table.append("| " + " | ".join(cells) + " |")
                        
                        content.append("\n".join(md_table))
                
    except Exception as e:
        content.append(f"PDF处理错误: {str(e)}")
    
    return "\n\n".join(content)

# 新增：将Markdown转换为Word文档
def convert_md_to_word(md_path, output_path):
    """
    将Markdown文件转换为Word文档
    :param md_path: Markdown文件路径
    :param output_path: 输出Word文件路径
    :return: 是否成功
    """
    try:
        # 使用pypandoc转换
        pypandoc.convert_file(
            md_path,
            'docx',
            outputfile=output_path,
            extra_args=['--reference-doc=reference.docx'] if os.path.exists('reference.docx') else []
        )
        return True
    except Exception as e:
        print(f"转换失败: {str(e)}")
        # 备用方法：使用python-docx手动转换
        try:
            # 读取Markdown文件
            with open(md_path, 'r', encoding='utf-8') as f:
                md_content = f.read()
            
            # 创建Word文档
            doc = docx.Document()
            
            # 处理标题和内容
            lines = md_content.split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    doc.add_paragraph()
                    continue
                    
                # 处理标题
                if line.startswith('#'):
                    level = 0
                    while line.startswith('#'):
                        level += 1
                        line = line[1:]
                    line = line.strip()
                    doc.add_heading(line, level=level)
                # 处理普通段落
                else:
                    p = doc.add_paragraph()
                    # 简单处理粗体和斜体
                    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', line)
                    for part in parts:
                        if part.startswith('**') and part.endswith('**'):
                            p.add_run(part[2:-2]).bold = True
                        elif part.startswith('*') and part.endswith('*'):
                            p.add_run(part[1:-1]).italic = True
                        elif part:
                            p.add_run(part)
            
            # 保存文档
            doc.save(output_path)
            return True
        except Exception as e2:
            print(f"备用方法也失败: {str(e2)}")
            return False

# 新增：将Markdown转换为PDF文档
def convert_md_to_pdf(md_path, output_path):
    """
    将Markdown文件转换为PDF
    :param md_path: Markdown文件路径
    :param output_path: 输出PDF路径
    :return: 是否成功
    """
    try:
        # 尝试使用pypandoc
        pypandoc.convert_file(
            md_path,
            'pdf',
            outputfile=output_path
        )
        return True
    except Exception as e:
        print(f"pypandoc转换失败: {str(e)}")
        
        # 备用方法：使用reportlab手动转换
        try:
            # 读取Markdown文件
            with open(md_path, 'r', encoding='utf-8') as f:
                md_content = f.read()
            
            # 将Markdown转换为HTML
            html = markdown.markdown(md_content, extensions=['tables', 'fenced_code'])
            
            # 创建PDF文档
            doc = SimpleDocTemplate(output_path, pagesize=letter)
            styles = getSampleStyleSheet()
            
            # 自定义标题样式
            styles.add(ParagraphStyle(name='Heading1', fontSize=16, bold=True))
            styles.add(ParagraphStyle(name='Heading2', fontSize=14, bold=True))
            styles.add(ParagraphStyle(name='Heading3', fontSize=12, bold=True))
            
            # 解析HTML并添加到PDF
            flowables = []
            
            # 简单解析HTML (这是一个非常基础的实现，真实情况需要更复杂的HTML解析)
            lines = html.split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    flowables.append(Spacer(1, 12))
                    continue
                
                # 处理标题
                if line.startswith('<h1>'):
                    text = line.replace('<h1>', '').replace('</h1>', '')
                    flowables.append(Paragraph(text, styles['Heading1']))
                elif line.startswith('<h2>'):
                    text = line.replace('<h2>', '').replace('</h2>', '')
                    flowables.append(Paragraph(text, styles['Heading2']))
                elif line.startswith('<h3>'):
                    text = line.replace('<h3>', '').replace('</h3>', '')
                    flowables.append(Paragraph(text, styles['Heading3']))
                # 处理段落
                elif not (line.startswith('<table>') or line.startswith('<tr>') or line.startswith('<td>')):
                    # 移除其他HTML标签
                    text = re.sub(r'<[^>]*>', '', line)
                    flowables.append(Paragraph(text, styles['Normal']))
            
            # 构建PDF
            doc.build(flowables)
            return True
        except Exception as e2:
            print(f"备用方法也失败: {str(e2)}")
            return False

class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.file_list = []
        self.init_ui()
        
    def create_button(self, text, callback=None, height=None, style=None):
        """创建统一样式的按钮"""
        btn = QPushButton(text)
        if callback:
            btn.clicked.connect(callback)
        if height:
            btn.setFixedHeight(height)
        if style:
            btn.setStyleSheet(style)
        return btn
        
    def init_ui(self):
        self.setWindowTitle("文档转换工具")
        self.setMinimumSize(800, 600)
        
        # 创建选项卡
        tabs = QTabWidget()
        
        # 创建"转Markdown"选项卡
        to_md_tab = QWidget()
        tabs.addTab(to_md_tab, "转换为Markdown")
        
        # 创建"从Markdown转换"选项卡
        from_md_tab = QWidget()
        tabs.addTab(from_md_tab, "从Markdown转换")
        
        # 设置"转Markdown"选项卡布局
        self.setup_to_md_tab(to_md_tab)
        
        # 设置"从Markdown转换"选项卡布局
        self.setup_from_md_tab(from_md_tab)
        
        # 设置主布局
        main_layout = QVBoxLayout()
        main_layout.addWidget(tabs)
        
        # 设置中心部件
        central_widget = QWidget()
        central_widget.setLayout(main_layout)
        self.setCentralWidget(central_widget)
        
        self.show()
    
    def setup_to_md_tab(self, tab):
        """设置转为Markdown选项卡的界面"""
        layout = QVBoxLayout(tab)
        
        # 标题
        title_label = QLabel("Word/PDF文档转Markdown工具")
        title_font = QFont()
        title_font.setPointSize(16)
        title_font.setBold(True)
        title_label.setFont(title_font)
        title_label.setAlignment(Qt.AlignCenter)
        layout.addWidget(title_label)
        
        # 文件类型选择
        filetype_group = QGroupBox("文件类型")
        filetype_layout = QVBoxLayout()
        self.to_md_filetype_group = QButtonGroup(self)
        
        self.word_type_radio = QRadioButton("Word文档 (*.docx, *.doc)")
        self.pdf_type_radio = QRadioButton("PDF文档 (*.pdf)")
        
        self.word_type_radio.setChecked(True)
        self.to_md_filetype_group.addButton(self.word_type_radio, 1)
        self.to_md_filetype_group.addButton(self.pdf_type_radio, 2)
        
        filetype_layout.addWidget(self.word_type_radio)
        filetype_layout.addWidget(self.pdf_type_radio)
        filetype_group.setLayout(filetype_layout)
        layout.addWidget(filetype_group)
        
        # 文件选择
        file_layout = QHBoxLayout()
        file_layout.addWidget(QLabel("选择文档:"))
        self.to_md_browse_btn = self.create_button("浏览...", self.browse_to_md_files)
        file_layout.addWidget(self.to_md_browse_btn)
        layout.addLayout(file_layout)
        
        # 文件列表
        self.to_md_file_list = QListWidget()
        layout.addWidget(self.to_md_file_list)
        
        # 转换模式选择 - 只对Word有效
        self.to_md_mode_group = QGroupBox("Word转换模式")
        mode_layout = QVBoxLayout()
        self.mode_button_group = QButtonGroup(self)
        
        self.simple_mode_radio = QRadioButton("简单转换 (一个Word文档转为一个Markdown文件)")
        self.sections_mode_radio = QRadioButton("分割转换 (按一级标题将文档分割为多个Markdown文件)")
        
        self.simple_mode_radio.setChecked(True)
        self.mode_button_group.addButton(self.simple_mode_radio, 1)
        self.mode_button_group.addButton(self.sections_mode_radio, 2)
        
        mode_layout.addWidget(self.simple_mode_radio)
        mode_layout.addWidget(self.sections_mode_radio)
        self.to_md_mode_group.setLayout(mode_layout)
        layout.addWidget(self.to_md_mode_group)
        
        # 合并选项
        self.to_md_merge_checkbox = QCheckBox("将多个文档合并为一个Markdown文件")
        layout.addWidget(self.to_md_merge_checkbox)
        
        # 连接文件类型切换事件
        self.word_type_radio.toggled.connect(self.toggle_to_md_mode_options)
        self.pdf_type_radio.toggled.connect(self.toggle_to_md_mode_options)
        
        # 输出目录
        dir_layout = QHBoxLayout()
        dir_layout.addWidget(QLabel("输出目录:"))
        self.to_md_dir_path = QLabel()
        
        # 设置默认输出目录
        default_download_dir = os.path.join(os.path.expanduser("~"), "Downloads")
        self.to_md_dir_path.setText(default_download_dir if os.path.exists(default_download_dir) else os.getcwd())
        self.to_md_dir_path.setStyleSheet("background-color: #f0f0f0; padding: 5px; border-radius: 3px;")
        
        self.to_md_browse_dir_btn = self.create_button("更改...", self.browse_to_md_directory)
        dir_layout.addWidget(self.to_md_dir_path, 1)
        dir_layout.addWidget(self.to_md_browse_dir_btn)
        layout.addLayout(dir_layout)
        
        # 转换按钮
        self.to_md_convert_btn = self.create_button(
            "开始转换", 
            self.start_to_md_conversion, 
            height=40, 
            style="background-color: #4CAF50; color: white; font-weight: bold;"
        )
        layout.addWidget(self.to_md_convert_btn)
        
        # 进度显示
        self.to_md_progress_label = QLabel("就绪")
        layout.addWidget(self.to_md_progress_label)
        
        self.to_md_progress_bar = QProgressBar()
        self.to_md_progress_bar.setValue(0)
        layout.addWidget(self.to_md_progress_bar)
        
        # 日志区域
        self.to_md_log_area = QTextEdit()
        self.to_md_log_area.setReadOnly(True)
        layout.addWidget(self.to_md_log_area)
    
    def setup_from_md_tab(self, tab):
        """设置从Markdown转换选项卡的界面"""
        layout = QVBoxLayout(tab)
        
        # 标题
        title_label = QLabel("Markdown转Word/PDF工具")
        title_font = QFont()
        title_font.setPointSize(16)
        title_font.setBold(True)
        title_label.setFont(title_font)
        title_label.setAlignment(Qt.AlignCenter)
        layout.addWidget(title_label)
        
        # 目标类型选择
        target_group = QGroupBox("目标格式")
        target_layout = QVBoxLayout()
        self.from_md_target_group = QButtonGroup(self)
        
        self.target_word_radio = QRadioButton("Word文档 (*.docx)")
        self.target_pdf_radio = QRadioButton("PDF文档 (*.pdf)")
        
        self.target_word_radio.setChecked(True)
        self.from_md_target_group.addButton(self.target_word_radio, 1)
        self.from_md_target_group.addButton(self.target_pdf_radio, 2)
        
        target_layout.addWidget(self.target_word_radio)
        target_layout.addWidget(self.target_pdf_radio)
        target_group.setLayout(target_layout)
        layout.addWidget(target_group)
        
        # 文件选择
        file_layout = QHBoxLayout()
        file_layout.addWidget(QLabel("选择Markdown文件:"))
        self.from_md_browse_btn = self.create_button("浏览...", self.browse_from_md_files)
        file_layout.addWidget(self.from_md_browse_btn)
        layout.addLayout(file_layout)
        
        # 文件列表
        self.from_md_file_list = QListWidget()
        layout.addWidget(self.from_md_file_list)
        
        # 合并选项
        self.from_md_merge_checkbox = QCheckBox("将多个Markdown文件合并为一个输出文件")
        layout.addWidget(self.from_md_merge_checkbox)
        
        # 输出目录
        dir_layout = QHBoxLayout()
        dir_layout.addWidget(QLabel("输出目录:"))
        self.from_md_dir_path = QLabel()
        
        # 设置默认输出目录
        default_download_dir = os.path.join(os.path.expanduser("~"), "Downloads")
        self.from_md_dir_path.setText(default_download_dir if os.path.exists(default_download_dir) else os.getcwd())
        self.from_md_dir_path.setStyleSheet("background-color: #f0f0f0; padding: 5px; border-radius: 3px;")
        
        self.from_md_browse_dir_btn = self.create_button("更改...", self.browse_from_md_directory)
        dir_layout.addWidget(self.from_md_dir_path, 1)
        dir_layout.addWidget(self.from_md_browse_dir_btn)
        layout.addLayout(dir_layout)
        
        # 转换按钮
        self.from_md_convert_btn = self.create_button(
            "开始转换", 
            self.start_from_md_conversion, 
            height=40, 
            style="background-color: #4CAF50; color: white; font-weight: bold;"
        )
        layout.addWidget(self.from_md_convert_btn)
        
        # 进度显示
        self.from_md_progress_label = QLabel("就绪")
        layout.addWidget(self.from_md_progress_label)
        
        self.from_md_progress_bar = QProgressBar()
        self.from_md_progress_bar.setValue(0)
        layout.addWidget(self.from_md_progress_bar)
        
        # 日志区域
        self.from_md_log_area = QTextEdit()
        self.from_md_log_area.setReadOnly(True)
        layout.addWidget(self.from_md_log_area)
    
    def toggle_to_md_mode_options(self):
        """根据选择的文件类型启用/禁用Word特有选项"""
        is_word = self.word_type_radio.isChecked()
        self.to_md_mode_group.setEnabled(is_word)
    
    def browse_to_md_files(self):
        """选择要转换为Markdown的文件"""
        if self.word_type_radio.isChecked():
            file_filter = "Word文档 (*.docx *.doc)"
        else:
            file_filter = "PDF文档 (*.pdf)"
            
        file_paths, _ = QFileDialog.getOpenFileNames(
            self, "选择文档", "", file_filter)
            
        if file_paths:
            self.to_md_file_list.clear()
            for file_path in file_paths:
                self.to_md_file_list.addItem(os.path.basename(file_path))
            self.to_md_log_area.append(f"已选择 {len(file_paths)} 个文件")
            self.to_md_file_paths = file_paths
    
    def browse_from_md_files(self):
        """选择要转换的Markdown文件"""
        file_paths, _ = QFileDialog.getOpenFileNames(
            self, "选择Markdown文件", "", "Markdown文件 (*.md)")
            
        if file_paths:
            self.from_md_file_list.clear()
            for file_path in file_paths:
                self.from_md_file_list.addItem(os.path.basename(file_path))
            self.from_md_log_area.append(f"已选择 {len(file_paths)} 个Markdown文件")
            self.from_md_file_paths = file_paths
    
    def browse_to_md_directory(self):
        """选择转Markdown的输出目录"""
        if dir_path := QFileDialog.getExistingDirectory(self, "选择输出目录", self.to_md_dir_path.text()):
            self.to_md_dir_path.setText(dir_path)
            self.to_md_log_area.append(f"已选择输出目录: {dir_path}")
    
    def browse_from_md_directory(self):
        """选择从Markdown转换的输出目录"""
        if dir_path := QFileDialog.getExistingDirectory(self, "选择输出目录", self.from_md_dir_path.text()):
            self.from_md_dir_path.setText(dir_path)
            self.from_md_log_area.append(f"已选择输出目录: {dir_path}")
    
    def start_to_md_conversion(self):
        """开始将文档转换为Markdown"""
        if not hasattr(self, 'to_md_file_paths') or not self.to_md_file_paths:
            QMessageBox.warning(self, "警告", "请先选择文档！")
            return
        
        # 禁用按钮，防止重复点击
        self.toggle_to_md_controls(False)
        
        # 获取文件类型
        file_type = 'pdf' if self.pdf_type_radio.isChecked() else 'word'
        
        # 获取合并选项
        merge_output = self.to_md_merge_checkbox.isChecked()
        
        # 获取转换模式 (仅Word文档需要考虑模式)
        if file_type == 'word':
            mode = 'simple' if self.simple_mode_radio.isChecked() else 'sections'
            mode_text = "简单模式" if mode == 'simple' else "分割模式"
            self.to_md_log_area.append(f"开始转换Word文档，使用{mode_text}，共 {len(self.to_md_file_paths)} 个文件...")
        else:
            mode = 'simple'  # PDF默认使用简单模式
            self.to_md_log_area.append(f"开始转换PDF文档，共 {len(self.to_md_file_paths)} 个文件...")
        
        if merge_output:
            self.to_md_log_area.append("文档将被合并为一个Markdown文件")
            
        # 启动转换线程
        self.to_md_thread = ToMarkdownThread(
            self.to_md_file_paths, 
            self.to_md_dir_path.text(), 
            mode, 
            merge_output, 
            file_type
        )
        self.to_md_thread.update_progress.connect(self.update_to_md_progress)
        self.to_md_thread.finished.connect(self.to_md_conversion_finished)
        self.to_md_thread.file_progress.connect(self.update_to_md_file_progress)
        self.to_md_thread.start()
    
    def start_from_md_conversion(self):
        """开始将Markdown转换为Word/PDF"""
        if not hasattr(self, 'from_md_file_paths') or not self.from_md_file_paths:
            QMessageBox.warning(self, "警告", "请先选择Markdown文件！")
            return
        
        # 禁用按钮，防止重复点击
        self.toggle_from_md_controls(False)
        
        # 获取目标格式
        target_format = 'pdf' if self.target_pdf_radio.isChecked() else 'word'
        
        # 获取合并选项
        merge_output = self.from_md_merge_checkbox.isChecked()
        
        # 日志输出
        format_name = "Word" if target_format == 'word' else "PDF"
        self.from_md_log_area.append(f"开始转换Markdown文件为{format_name}，共 {len(self.from_md_file_paths)} 个文件...")
        
        if merge_output and len(self.from_md_file_paths) > 1:
            self.from_md_log_area.append("Markdown文件将被合并为一个输出文件")
        
        # 启动转换线程
        self.from_md_thread = FromMarkdownThread(
            self.from_md_file_paths,
            self.from_md_dir_path.text(),
            target_format,
            merge_output
        )
        self.from_md_thread.update_progress.connect(self.update_from_md_progress)
        self.from_md_thread.finished.connect(self.from_md_conversion_finished)
        self.from_md_thread.file_progress.connect(self.update_from_md_file_progress)
        self.from_md_thread.start()
    
    def toggle_to_md_controls(self, enabled=True):
        """启用或禁用转Markdown选项卡的UI控件"""
        self.to_md_convert_btn.setEnabled(enabled)
        self.to_md_browse_btn.setEnabled(enabled)
        self.to_md_browse_dir_btn.setEnabled(enabled)
        self.word_type_radio.setEnabled(enabled)
        self.pdf_type_radio.setEnabled(enabled)
        self.simple_mode_radio.setEnabled(enabled and self.word_type_radio.isChecked())
        self.sections_mode_radio.setEnabled(enabled and self.word_type_radio.isChecked())
        self.to_md_merge_checkbox.setEnabled(enabled)
    
    def toggle_from_md_controls(self, enabled=True):
        """启用或禁用从Markdown转换选项卡的UI控件"""
        self.from_md_convert_btn.setEnabled(enabled)
        self.from_md_browse_btn.setEnabled(enabled)
        self.from_md_browse_dir_btn.setEnabled(enabled)
        self.target_word_radio.setEnabled(enabled)
        self.target_pdf_radio.setEnabled(enabled)
        self.from_md_merge_checkbox.setEnabled(enabled)
    
    def update_to_md_progress(self, value, message):
        """更新转Markdown选项卡的进度"""
        self.to_md_progress_bar.setValue(value)
        self.to_md_log_area.append(message)
    
    def update_from_md_progress(self, value, message):
        """更新从Markdown转换选项卡的进度"""
        self.from_md_progress_bar.setValue(value)
        self.from_md_log_area.append(message)
    
    def update_to_md_file_progress(self, current, total):
        """更新转Markdown选项卡的文件进度"""
        self.to_md_progress_label.setText(f"处理文件 {current}/{total}")
    
    def update_from_md_file_progress(self, current, total):
        """更新从Markdown转换选项卡的文件进度"""
        self.from_md_progress_label.setText(f"处理文件 {current}/{total}")
    
    def to_md_conversion_finished(self, success, message):
        """转Markdown完成时的回调"""
        if success:
            self.to_md_log_area.append(message)
            QMessageBox.information(self, "转换完成", message)
        else:
            self.to_md_log_area.append(f"错误: {message}")
            QMessageBox.critical(self, "转换失败", message)
        
        self.toggle_to_md_controls(True)
        self.to_md_progress_label.setText("就绪")
    
    def from_md_conversion_finished(self, success, message):
        """从Markdown转换完成时的回调"""
        if success:
            self.from_md_log_area.append(message)
            QMessageBox.information(self, "转换完成", message)
        else:
            self.from_md_log_area.append(f"错误: {message}")
            QMessageBox.critical(self, "转换失败", message)
        
        self.toggle_from_md_controls(True)
        self.from_md_progress_label.setText("就绪")

if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = MainWindow()
    sys.exit(app.exec_()) 